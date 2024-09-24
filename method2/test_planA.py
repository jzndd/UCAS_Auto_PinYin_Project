from docx import Document
from docx.shared import Pt
import json
from pypinyin import pinyin, Style

file_path_1 = 'polyphone.json'
file_path_2 = 'nlp_test.docx'
output_doc_path = 'nlp_pinyin.docx'

# 打开文件并加载多音字数据
with open(file_path_1, 'r', encoding='utf-8') as file:
    polyphone_data = json.load(file)


# 定义读取文档函数
def read_from_word_file(file_path):
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]


# 定义一个函数来处理输入文本，并检查其中的多音字，并添加注音
def add_pinyin_to_polyphone_words(paragraphs, polyphone_data):
    polyphonic_chars = {item['char'] for item in polyphone_data}
    output_paragraphs = []

    for paragraph in paragraphs:
        pinyin_paragraph = pinyin(paragraph, style=Style.TONE, errors='default')
        print(f"Original text: {paragraph}")
        print(f"Pinyin result: {pinyin_paragraph}")
        output_paragraph = ""
        pinyin_index = 0

        for char in paragraph:
            if char in polyphonic_chars:
                while pinyin_index < len(pinyin_paragraph) and not pinyin_paragraph[pinyin_index]:
                    pinyin_index += 1
                if pinyin_index < len(pinyin_paragraph):
                    output_paragraph += char + '(' + '/'.join(pinyin_paragraph[pinyin_index]) + ')'  # 处理多音字，显示多个拼音
                    pinyin_index += 1
                else:
                    output_paragraph += char  # 防止索引越界
            else:
                output_paragraph += char
                if pinyin_index < len(pinyin_paragraph) and pinyin_paragraph[pinyin_index]:
                    pinyin_index += 1

        output_paragraphs.append(output_paragraph)
    return output_paragraphs


# 读取Word文档内容
input_paragraphs_from_word = read_from_word_file(file_path_2)

# 处理文本并添加注音
output_paragraphs_with_pinyin = add_pinyin_to_polyphone_words(input_paragraphs_from_word, polyphone_data)

# 创建新的Word文档并写入处理后的文本
output_doc = Document()
for paragraph in output_paragraphs_with_pinyin:
    # 添加段落并设置字体和大小
    p = output_doc.add_paragraph(paragraph)
    p.paragraph_format.first_line_indent = Pt(21)  # 设置首行缩进两个字符
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)  # 小四号字体大小为12磅

print(output_paragraphs_with_pinyin)
output_doc.save(output_doc_path)
