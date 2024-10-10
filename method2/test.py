import torch
from torch import nn
from torch.autograd import Variable
import docx
import json
import re
from train import DisambiguationLSTM, make_sequence
import argparse
from nlp_nan import pinyin_nan

import jieba

# 
if __name__ == '__main__':

    parser = argparse.ArgumentParser()
    parser.add_argument('--scale', type=str, default='v4', help='v1 是llm构建+人工清洗的数据集，v2 是人民日报文本+pypinyin自动构建的数据集, v3是群聊 csv 提供的数据集, v4 是群聊 sent 和 lb 数据集,v4_l 是更大的数据集和网络结构')
    parser.add_argument('--input', type=str, default='data/nlp_test.docx', help='输入文件名')
    parser.add_argument('--output', type=str, default='data/nlp_test_output.docx', help='输出文件名')
    parser.add_argument('--polyphone', type=str, default='data/polyphone.json', help='多音字路径')
    parser.add_argument('--rare_char', type=str, default= 'data/rare_char.json', help='生僻字路径')
    parser.add_argument('--level', type=int, default= 3000, help='设置生僻字级别')
    parser.add_argument("--use_jieba", action="store_true", help="是否使用 jieba ")
    args = parser.parse_args()

    if args.scale == 'v1':
        train_data_file = 'data/train_data.json'
        model_file = 'data/disambiguation_models.pth'
        args.output = 'data/nlp_test_output.docx'
    elif args.scale == 'v2':
        train_data_file = 'data/train_data_v2.json'
        model_file = 'data/disambiguation_models_v2.pth'
        args.output = 'data/nlp_test_output_v2.docx'
    elif args.scale == 'v3':
        train_data_file = 'data/train_data_v3.json'
        model_file = 'data/disambiguation_models_v3.pth'
        args.output = 'data/nlp_test_output_v3.docx'
    elif args.scale == 'v4':
        train_data_file = 'data/train_data_v4.json'
        model_file = 'data/disambiguation_models_v4.pth'
        args.output = 'data/nlp_test_output_v4.docx'
    else:
        raise ValueError('scale 参数只能是 v1 或 v2')
    
    # 规则推理 ： 着，了，的，地 做单个字时发声是确定地
    rule_word = {"着":"zhe", "了":"le", "的":"de", "地":"de"}

    standard_dict = json.load(open('data/standard_pron.json', 'r', encoding='utf-8'))

    poly_pronunciation = json.load(open('data/polyphone_converted_data.json', 'r', encoding='utf-8'))
    poly_dict = {entry['char']: entry['pinyin'] for entry in poly_pronunciation}

    # 读取训练数据
    with open(train_data_file, 'r', encoding='utf-8') as f:
        train_data = json.load(f)

    # 将每个字和多音字的注音编码
    word_to_idx = {}
    pron_to_idx = {}
    for word, examples in train_data.items():
        if word not in word_to_idx:
            word_to_idx[word] = len(word_to_idx)
        for example in examples:
            _, _, pron = example
            for ch in example[0]:  # 确保句子中的每个字都被编码
                if ch not in word_to_idx:
                    word_to_idx[ch] = len(word_to_idx)
            if pron not in pron_to_idx:
                pron_to_idx[pron] = len(pron_to_idx)

    # 加载模型
    models = nn.ModuleDict()
    for word in train_data.keys():
        models[word] = DisambiguationLSTM(len(word_to_idx) + 1, 100, 128, len(pron_to_idx))
    models.load_state_dict(torch.load(model_file))
    models.eval()  

    # 读取docx文件
    doc = docx.Document(args.input)  
    print("读取文档成功")

    # 定义句子分割的正则表达式（包括标点符号）
    sentence_pattern = re.compile(r'([^，。；：“”]+[，。；：“”]?)')

    output_doc = docx.Document()
    # 遍历文档中的段落
    for para in doc.paragraphs:
        text = para.text
        sentences = sentence_pattern.findall(text)  # 使用正则表达式分割句子并保留标点符号
        
        # 创建一个新的段落内容
        new_paragraph = ""
        
        for sentence in sentences:
            word_dict = {}
            sentence_jieba_cut = list(jieba.lcut(sentence))
            new_sentence = sentence  # 初始化为原始句子
            char_index = 0  # 用于记录当前处理的字符位置
            # visited_word = set()  # 用于记录已经处理过的词语
            # for word in train_data.keys():
            for i in range(len(sentence)):
                # if word in sentence and word not in visited_word:
                char_index += 1
                word = sentence[i]
                jieba_cut_cnt = 0
                if word in train_data.keys():
                    # 动态统计 word 出现的次数
                    if word not in word_dict:
                        word_dict[word] = 0
                    word_dict[word] += 1

                    for sub_sentence_jieba_cut in sentence_jieba_cut:
                        if word in sub_sentence_jieba_cut:
                            jieba_cut_cnt +=1
                            if jieba_cut_cnt == word_dict[word]:
                                break

                    # 规则推理 ： 着，了，的，地 做单个字时发声是确定的
                    if word in rule_word and len(sub_sentence_jieba_cut)==1:
                        predicted_pron = rule_word[word]
                    else:
                        # 模型推理
                        if args.use_jieba and  len(sub_sentence_jieba_cut)!=1 :
                            input_sentence = sub_sentence_jieba_cut
                        else:
                            input_sentence = sentence 

                        input_seq = make_sequence(input_sentence, word_to_idx).unsqueeze(0)  # 添加 batch 维度
                        with torch.no_grad():
                            output = models[word](input_seq)
                            pred_index = torch.max(output, 1)[1].item()  # 获取预测的拼音索引
                            predicted_pron = list(pron_to_idx.keys())[list(pron_to_idx.values()).index(pred_index)]
                        
                    # 在字符后添加拼音
                    new_sentence = new_sentence[:char_index] + f'({predicted_pron})' + new_sentence[char_index:]
                    char_index += len(predicted_pron) + 2

                elif word not in train_data.keys() and word in standard_dict.keys() and word in poly_dict.keys():\
                    # 在字符后添加拼音
                    new_sentence = new_sentence[:char_index] + f'({standard_dict[word]})' + new_sentence[char_index:]
                    char_index += len(standard_dict[word][0]) + 2

                elif word not in train_data.keys() and word in standard_dict.keys() and word not in poly_dict.keys():
                    # 在字符后添加拼音
                    new_sentence = new_sentence[:char_index] + f'({standard_dict[word][0]})' + new_sentence[char_index:]
                    char_index += len(standard_dict[word][0]) + 2

            new_paragraph += new_sentence  # 保留句子之间的空格
        
        output_doc.add_paragraph(new_paragraph)  # 保持段落结构
      
    # 保存结果到新的docx文件
    output_doc.save(args.output)
    # print("开始难字检测") 
    # pinyin_nan(input_file=args.output,output_doc_path=args.output,polyphone =args.polyphone,rare_char =args.rare_char,level = args.level)
    # print("拼音已添加并保存到 output.docx")
