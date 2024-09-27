from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from pypinyin import pinyin, Style

# 定义读取文档函数
def read_from_word_file(file_path):
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]


def get_pinyin(char,rare_char_data):
    for item in rare_char_data:
        if item['char'] == char:
            return item['pinyin']
    return None  
# 定义一个函数来处理输入文本，并检查其中的多音字与生僻字，并添加注音
def add_pinyin_to_polyphone_words(paragraphs, polyphone_data,rare_char_data):
    polyphonic_chars = {item['char'] for item in polyphone_data}
    rare_chars = {item['char'] for item in rare_char_data}
    output_paragraphs = []

    for paragraph in paragraphs:
        # pinyin_paragraph = pinyin(paragraph, style=Style.TONE)
        output_paragraph = ""
        char_index = 0
        for char in paragraph:
            # if char in polyphonic_chars :
            #     output_paragraph += char + '(' + pinyin_paragraph[char_index][0] + ')'
            if char in rare_chars:
                output_paragraph += char + '[' + get_pinyin(char,rare_char_data)[0] + ']'
            else:
                output_paragraph += char
            char_index += 1
        output_paragraphs.append(output_paragraph)
    return output_paragraphs

def pinyin_nan(input_file = 'method1/nlp_test.docx',output_doc_path = 'output_with_pinyin.docx',polyphone = 'data/polyphone.json',rare_char = 'data/rare_char.json',level = 3000):
    # 打开文件并加载多音字数据
    with open(polyphone, 'r', encoding='utf-8') as file:
        polyphone_data = json.load(file)
    with open(rare_char, 'r', encoding='utf-8') as file:
        rare_char_data = json.load(file)
        rare_char_data = rare_char_data[level:]
        
    # 读取Word文档内容
    input_paragraphs_from_word = read_from_word_file(input_file)

    # 处理文本并添加注音
    output_paragraphs_with_pinyin = add_pinyin_to_polyphone_words(input_paragraphs_from_word, polyphone_data,rare_char_data)

    # 创建新的Word文档并写入处理后的文本
    output_doc = Document()
    for paragraph in output_paragraphs_with_pinyin:
        # 添加段落并设置字体和大小
        p = output_doc.add_paragraph(paragraph)
        p.paragraph_format.first_line_indent = Pt(21)  # 设置首行缩进两个字符
        for run in p.runs:
            run.font.name = '宋体'
            run.font.size = Pt(12)  # 小四号字体大小为12磅
    output_doc.save(output_doc_path)

# pinyin_nan()