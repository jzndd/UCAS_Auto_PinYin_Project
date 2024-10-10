import re
import torch
from torch import nn
from torch.autograd import Variable
import docx
import json
import re
from train import DisambiguationLSTM, make_sequence
import argparse

import jieba 

import sys

sys.path.append('..')

# 提取注音
def extract_pinyin(text):
    pattern = r'\((.*?)\)'  # 正则表达式模式，匹配括号内的内容
    return re.findall(pattern, text)  # 返回所有匹配的注音列表

# 声调转换成数字
def convert_tone_to_number(pinyin_list):
    # 声调映射表
    tone_map = {
        'ā': 'a1', 'á': 'a2', 'ǎ': 'a3', 'à': 'a4',
        'ē': 'e1', 'é': 'e2', 'ě': 'e3', 'è': 'e4',
        'ī': 'i1', 'í': 'i2', 'ǐ': 'i3', 'ì': 'i4',
        'ō': 'o1', 'ó': 'o2', 'ǒ': 'o3', 'ò': 'o4',
        'ū': 'u1', 'ú': 'u2', 'ǔ': 'u3', 'ù': 'u4'
    }
    result = []
    for p in pinyin_list:
        # 分离声母和韵母
        tone = ''
        new_pinyin = ''
        for c in p:
            if c in tone_map:
                # 如果是声调字符，提取声调数字
                tone = tone_map[c][-1]
                # 替换为无声调的拼音
                new_pinyin += tone_map[c][:-1]
            else:
                new_pinyin += c
        # 将声调数字移动到拼音的最后方
        result.append(new_pinyin + tone)
    
    return result

# 计算最长公共子序列长度
def longest_common_subsequence_length(X, Y):
    m = len(X)
    n = len(Y)
    dp = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(m + 1):
        for j in range(n + 1):
            if i == 0 or j == 0:
                dp[i][j] = 0
            elif X[i - 1] == Y[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
            else:
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])

    return dp[m][n]

if __name__ == '__main__':

    parser = argparse.ArgumentParser()
    parser.add_argument('--scale', type=str, default='v4', help='v1 是llm构建+人工清洗的数据集，v2 是人民日报文本+pypinyin自动构建的数据集, v3是群聊 csv 提供的数据集, v4 是群聊 sent 和 lb 数据集,v4_l 是更大的数据集和网络结构')
    parser.add_argument('--input', type=str, default='data/eval.docx', help='输入文件名')
    parser.add_argument('--gt', type=str, default='data/eval_gt.docx', help='ground truhe 文件名')
    parser.add_argument('--polyphone', type=str, default='data/polyphone.json', help='多音字路径')
    parser.add_argument('--rare_char', type=str, default= 'data/rare_char.json', help='生僻字路径')
    parser.add_argument('--level', type=int, default= 3000, help='设置生僻字级别')
    parser.add_argument("--use_jieba", action="store_true", help="是否使用 jieba ")
    args = parser.parse_args()

    # ------------------------------------- load model -------------------------------------
    print("use jieba is :",args.use_jieba)
    # 规则推理 ： 着，了，的，地 做单个字时发声是确定的
    rule_word = {"着":"zhe", "了":"le", "的":"de", "地":"de"}

    standard_dict = json.load(open('data/standard_pron.json', 'r', encoding='utf-8'))

    poly_pronunciation = json.load(open('data/polyphone_converted_data.json', 'r', encoding='utf-8'))
    poly_dict = {entry['char']: entry['pinyin'] for entry in poly_pronunciation}

    if args.scale == 'v1':
        train_data_file = 'data/train_data.json'
        model_file = 'data/disambiguation_models.pth'
    elif args.scale == 'v2':
        train_data_file = 'data/train_data_v2.json'
        model_file = 'data/disambiguation_models_v2.pth'
    elif args.scale == 'v3':
        train_data_file = 'data/train_data_v3.json'
        model_file = 'data/disambiguation_models_v3.pth'
    elif args.scale == 'v4':
        train_data_file = 'data/train_data_v4.json'
        model_file = 'data/disambiguation_models_v4.pth'
    else:
        raise ValueError('scale 参数只能是 v1 或 v2')

    # 读取训练数据
    with open(train_data_file, 'r', encoding='utf-8') as f:
        train_data = json.load(f)

    # 将每个字和多音字的注音编码
    word_to_idx = {}
    pron_to_idx = {}
    for word, examples in train_data.items():
        if word not in word_to_idx:
            word_to_idx[word] = len(word_to_idx)
        for example in examples:
            _, _, pron = example
            for ch in example[0]:  # 确保句子中的每个字都被编码
                if ch not in word_to_idx:
                    word_to_idx[ch] = len(word_to_idx)
            if pron not in pron_to_idx:
                pron_to_idx[pron] = len(pron_to_idx)

    # 加载模型
    models = nn.ModuleDict()
    for word in train_data.keys():
        models[word] = DisambiguationLSTM(len(word_to_idx) + 1, 100, 128, len(pron_to_idx))
    models.load_state_dict(torch.load(model_file))
    models.eval()  


    # ------------------------------------- load gt -------------------------------------

    # 读取docx文件
    doc = docx.Document(args.input)  
    gt_doc = docx.Document(args.gt)
    gt_pron_list = []
    # 遍历gt文档中的段落
    for para in gt_doc.paragraphs:
        text = para.text
        gt_pron_list += extract_pinyin(text)  

    # 定义句子分割的正则表达式（包括标点符号）
    sentence_pattern = re.compile(r'([^，。；：“”]+[，。；：“”]?)')

    predicted_pron_list = []
    # 遍历文档中的段落
    for para in doc.paragraphs:
        text = para.text
        sentences = sentence_pattern.findall(text)  # 使用正则表达式分割句子并保留标点符号
        
        for sentence in sentences:
            # visited_word = set()  # 用于记录已经处理过的词语
            sentence_jieba_cut = jieba.lcut(sentence)
            word_dict = {}
            for word in sentence:

                jieba_cut_cnt = 0

                if word in train_data.keys():

                    # 动态统计 word 出现的次数
                    if word not in word_dict:
                        word_dict[word] = 0
                    word_dict[word] += 1

                    for sub_sentence_jieba_cut in sentence_jieba_cut:
                        if word in sub_sentence_jieba_cut:
                            jieba_cut_cnt +=1
                            if jieba_cut_cnt == word_dict[word]:
                                break

                    # 规则推理 ： 着，了，的，地 做单个字时发声是确定的
                    if word in rule_word and len(sub_sentence_jieba_cut)==1:
                        predicted_pron = rule_word[word]
                    else:
                        # 模型推理
                        if args.use_jieba and  len(sub_sentence_jieba_cut)!=1 :
                            input_sentence = sub_sentence_jieba_cut
                        else:
                            input_sentence = sentence  

                        # 使用模型进行推理
                        input_seq = make_sequence(input_sentence, word_to_idx).unsqueeze(0)  # 添加 batch 维度
                        with torch.no_grad():
                            output = models[word](input_seq)
                            pred_index = torch.max(output, 1)[1].item()  # 获取预测的拼音索引
                            predicted_pron = list(pron_to_idx.keys())[list(pron_to_idx.values()).index(pred_index)]
                        
                    predicted_pron_list.append(predicted_pron)

                elif word not in train_data.keys() and word in standard_dict.keys() and word in poly_dict.keys():
                    predicted_pron_list.append(standard_dict[word])

                elif word not in train_data.keys() and word not in standard_dict.keys() and word in poly_dict.keys():
                    predicted_pron_list.append(poly_dict[word][0])

    gt_pinyin_list = convert_tone_to_number(gt_pron_list)
    predicted_pron_list = convert_tone_to_number(predicted_pron_list)

    # print(gt_pron_list)

    # print("-----------------------------------------------------")

    # print(predicted_pron_list)

    lcs_length = longest_common_subsequence_length(gt_pinyin_list, predicted_pron_list)

    acc = lcs_length / len(gt_pinyin_list)

    print(acc)