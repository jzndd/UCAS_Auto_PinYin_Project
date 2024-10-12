import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from pypinyin import pinyin, Style
import docx
import sys
import re

sys.path.append("..")

# 计算最长公共子序列长度
def longest_common_subsequence_length(X, Y):
    m = len(X)
    n = len(Y)
    dp = [[0] * (n + 1) for _ in range(m + 1)]

    for i in range(m + 1):
        for j in range(n + 1):
            if i == 0 or j == 0:
                dp[i][j] = 0
            elif X[i - 1] == Y[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
            else:
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])

    return dp[m][n]

# 提取注音
def extract_pinyin(text):
    pattern = r'\((.*?)\)'  # 正则表达式模式，匹配括号内的内容
    return re.findall(pattern, text)  # 返回所有匹配的注音列表

def convert_tone_to_number(pinyin_list):
    # 声调映射表
    tone_map = {
        'ā': 'a1', 'á': 'a2', 'ǎ': 'a3', 'à': 'a4',
        'ē': 'e1', 'é': 'e2', 'ě': 'e3', 'è': 'e4',
        'ī': 'i1', 'í': 'i2', 'ǐ': 'i3', 'ì': 'i4',
        'ō': 'o1', 'ó': 'o2', 'ǒ': 'o3', 'ò': 'o4',
        'ū': 'u1', 'ú': 'u2', 'ǔ': 'u3', 'ù': 'u4'
    }
    result = []
    for p in pinyin_list:
        # 分离声母和韵母
        tone = ''
        new_pinyin = ''
        for c in p:
            if c in tone_map:
                # 如果是声调字符，提取声调数字
                tone = tone_map[c][-1]
                # 替换为无声调的拼音
                new_pinyin += tone_map[c][:-1]
            else:
                new_pinyin += c
        # 将声调数字移动到拼音的最后方
        result.append(new_pinyin + tone)
    
    return result

# 定义读取文档函数
def read_from_word_file(file_path):
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]

# 定义一个函数来处理输入文本，并检查其中的多音字与生僻字，并添加注音
def add_pinyin_to_polyphone_words(paragraphs, polyphone_data):
    predicted_pron_list = []

    for paragraph in paragraphs:
        pinyin_paragraph = pinyin(paragraph, style=Style.TONE)
        for i, char in enumerate(paragraph):
            if char in polyphone_data.keys():
                predicted_pron_list.append(pinyin_paragraph[i][0])

    return predicted_pron_list

if __name__ == '__main__':

    parser = argparse.ArgumentParser()
    parser.add_argument('--benchmark', type=str, default='v2', help='v1 是正常文本，v2是极端文本')
    parser.add_argument('--polyphone', type=str, default='data/polyphone.json', help='多音字路径')
    args = parser.parse_args()

    poly_pronunciation = json.load(open(args.polyphone, 'r', encoding='utf-8'))
    poly_dict = {entry['char']: entry['pinyin'] for entry in poly_pronunciation}

    if args.benchmark == 'v1':
        args.input = 'data/eval.docx'
        args.gt = 'data/eval_gt.docx'
    elif args.benchmark == 'v2':
        args.input = 'data/eval2.docx'
        args.gt = 'data/eval2_gt.docx'
    else:
        raise ValueError('benchmark 参数只能是 v1 或 v2')

    gt_doc = docx.Document(args.gt)
    gt_pron_list = []

    # 遍历gt文档中的段落
    for para in gt_doc.paragraphs:
        text = para.text
        gt_pron_list += extract_pinyin(text) 

    # 读取Word文档内容
    input_paragraphs_from_word = read_from_word_file(args.input)

    # 注音
    predicted_pron_list = add_pinyin_to_polyphone_words(input_paragraphs_from_word, poly_dict)

    gt_pinyin_list = convert_tone_to_number(gt_pron_list)
    predicted_pron_list = convert_tone_to_number(predicted_pron_list)

    lcs_length = longest_common_subsequence_length(gt_pinyin_list, predicted_pron_list)

    acc = lcs_length / len(gt_pinyin_list)

    print(acc)
